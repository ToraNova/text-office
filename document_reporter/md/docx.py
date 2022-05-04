# renders a docx from markdown
# using mistletoe for markdown parsing and python-docx for docx rendering

import re
from urllib.parse import urlparse

from itertools import chain
from mistletoe import block_token, span_token
from mistletoe.base_renderer import BaseRenderer

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..util.docx_helper import format_paragraph, format_run, format_table, insert_hrule
from .format_tag import HorizontalRuleTag, PageBreakTag, ImageTag, ColorTag, BoldTag, ItalicTag, UnderlineTag, StrikethroughTag, FontTag, AlignBlockTag, TableStyleBlockTag, ParagraphStyleBlockTag

class Renderer(BaseRenderer):
    def __init__(self, out_path, docx_template=None, rel_root=None, *extras):
        '''
        renders a .docx on out_path, using docx_template as the template file
        rel_root determine the path to include relative path resources (e.g., images)
        '''
        self._suppress_ptag_stack = [False]
        super().__init__(*chain(([
            HorizontalRuleTag, PageBreakTag, ImageTag,
            ColorTag, BoldTag, ItalicTag, UnderlineTag, StrikethroughTag, FontTag,
            AlignBlockTag, TableStyleBlockTag, ParagraphStyleBlockTag
        ]), extras))
        self.out_path = out_path
        self.rel_root = rel_root
        self.docx_template = docx_template

    def render_inner(self, token):
        # inner rendering not supported on DocxRenderer
        # raise NotImplementedError('unsupported render_inner function called')
        for c in token.children:
            self.render(c)

    def render_line_break(self, token):
        self.runs[-1].add_break()

    def render_image(self, token):

        print(token)
        print(token.title)
        print(token.src)
        up = urlparse(token.src)
        if up.scheme == 'file':
            if up.path.startswith('/'):
                img_path = up.path
            else:
                raise NotImplementedError('todo')
        else:
            raise NotImplementedError(f"unsupported image resource: {up.scheme}")

        self.runs.append(self.paras[-1].add_run())
        # TODO: introduce way to alter width
        self.runs[-1].add_picture(img_path, width=Cm(14))

    def render_raw_text(self, token):
        # add run to last added paragraph
        self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_text(token.content)

    def render_document(self, token):
        # create document from template
        if self.docx_template is None:
            self.docx = Document()
        else:
            self.docx = Document(self.docx_template)

        # paragraph and run stack
        self.runs = []
        self.paras = self.docx.paragraphs.copy()
        self.tables = self.docx.tables.copy()
        self.render_inner(token)

        # save document
        self.docx.save(self.out_path)

    def populate_and_format_runs(self, token, **kwargs):
        # format_tag allows nested formatting <b><i>test</i></b>, italic and bold
        tos = len(self.runs) # top of stack
        self.render_inner(token)
        added = len(self.runs) - tos # new top of stack
        # apply format to all added elements during inner render
        for i in range(added):
            format_run(self.runs[-(i+1)], **kwargs)

    def populate_and_format_paras(self, token, **kwargs):
        tos = len(self.paras)
        self.render_inner(token)
        added = len(self.paras) - tos
        # apply format to all added elements during inner render
        for i in range(added):
            format_paragraph(self.paras[-(i+1)], **kwargs)

    def populate_and_format_tables(self, token, **kwargs):
        tos = len(self.tables) # top of stack
        self.render_inner(token)
        added = len(self.tables) - tos # new top of stack
        # add style to all added table within this block
        for i in range(added):
            format_table(self.tables[-(i+1)], **kwargs)

    def render_color_tag(self, token):
        self.populate_and_format_runs(token, color=RGBColor.from_string(token.format_value))

    def render_bold_tag(self, token):
        self.populate_and_format_runs(token, bold=True)

    def render_strong(self, token):
        self.populate_and_format_runs(token, bold=True)

    def render_italic_tag(self, token):
        self.populate_and_format_runs(token, italic=True)

    def render_emphasis(self, token):
        self.populate_and_format_runs(token, italic=True)

    def render_underline_tag(self, token):
        self.populate_and_format_runs(token, underline=True)

    def render_strikethrough_tag(self, token):
        self.populate_and_format_runs(token, strike=True)

    def render_strikethrough(self, token):
        self.populate_and_format_runs(token, strike=True)

    def render_font_tag(self, token):
        self.populate_and_format_runs(token, name=token.format_value)

    def render_heading(self, token):
        # assume that heading has no additional child
        self.paras.append(self.docx.add_heading(level=token.level).clear())
        self.render_inner(token)

    def render_paragraph(self, token):
        # for all span tokens
        self.paras.append(self.docx.add_paragraph().clear())
        self.render_inner(token)

    def render_align_block_tag(self, token):
        map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if token.format_value not in map:
            raise KeyError(f'no such alignment type: {token.format_value} (center, left, right, justify)')
        self.populate_and_format_paras(token, align=map[token.format_value])

    def render_horizontal_rule_tag(self, token):
        insert_hrule(self.paras[-1], token.format_value)

    def render_page_break_tag(self, token):
        self.docx.add_page_break()

    def populate_table(self, token, **kwargs):
        if len(token.children) < 1:
            return
        ncol = len(token.children[0].children)
        self.tables.append(self.docx.add_table(0, ncol))

        if hasattr(token, 'header') and token.header is not None:
            token.header.is_header = True
            self.render(token.header)

        for row in token.children:
            row.is_header = False
            self.render(row)

    def populate_row(self, token, **kwargs):
        row = self.tables[-1].add_row()
        for cidx, col in enumerate(token.children):
            col.docx_cell = row.cells[cidx]
            self.render(col)

    def render_table(self, token):
        self.populate_table(token)

    def render_table_row(self, token):
        self.populate_row(token)

    def render_table_cell(self, token):
        self.paras.append(token.docx_cell.paragraphs[0])
        self.render_inner(token)

    def render_table_style_block_tag(self, token):
        self.populate_and_format_tables(token, style=token.format_value_raw)

    def render_paragraph_style_block_tag(self, token):
        self.populate_and_format_paras(token, style=token.format_value_raw)
