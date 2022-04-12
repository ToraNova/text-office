# https://mistletoe-ebp.readthedocs.io/en/latest/using/develop.html?highlight=renderer#a-new-renderer
# https://github.com/miyuchina/mistletoe/blob/master/mistletoe/html_renderer.py

import re

from itertools import chain
from mistletoe import block_token, span_token
from mistletoe.base_renderer import BaseRenderer

from docx import Document
from docx.shared import RGBColor

from .docx_helper import format_paragraph, format_run
from .format_token import ColorSpan, BoldSpan, ItalicSpan

class Renderer(BaseRenderer):
    def __init__(self, out_path, docx_template=None, *extras):
        self._suppress_ptag_stack = [False]
        super().__init__(*chain(([ColorSpan, BoldSpan, ItalicSpan]), extras))
        self.out_path = out_path
        self.docx_template = docx_template

    def render_inner(self, token):
        # inner rendering not supported on DocxRenderer
        # raise NotImplementedError('unsupported render_inner function called')
        for c in token.children:
            self.render(c)

    def render_line_break(self, token):
        self.runs[-1].add_break()

    def render_raw_text(self, token):
        # add run to last added paragraph
        self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_text(token.content)

    def render_document(self, token):
        # create document from template
        if self.docx_template is None:
            self.docx = Document()
        else:
            self.docx = Document(self.docx_template)

        # paragraph and run stack
        self.paras = []
        self.runs = []
        self.render_inner(token)

        # save document
        self.docx.save(self.out_path)

    def render_color_span(self, token):
        # format_span allows nested formatting <b><i>test</i></b>, italic and bold
        tos = len(self.runs) # top of stack
        self.render_inner(token)
        added = len(self.runs) - tos # new top of stack
        # apply format to all added elements during inner render
        for i in range(added):
            format_run(self.runs[-(i+1)], color=RGBColor.from_string(token.format_value))

    def render_bold_span(self, token):
        tos = len(self.runs) # top of stack
        self.render_inner(token)
        added = len(self.runs) - tos # new top of stack
        # apply format to all added elements during inner render
        for i in range(added):
            format_run(self.runs[-(i+1)], bold=True)

    def render_italic_span(self, token):
        tos = len(self.runs) # top of stack
        self.render_inner(token)
        added = len(self.runs) - tos # new top of stack
        # apply format to all added elements during inner render
        for i in range(added):
            format_run(self.runs[-(i+1)], italic=True)

    def render_strong(self, token):
        self.render_inner(token)
        format_run(self.runs[-1], bold=True)

    def render_emphasis(self, token):
        self.render_inner(token)
        format_run(self.runs[-1], italic=True)

    def render_heading(self, token):
        # assume that heading has no additional child
        self.paras.append(self.docx.add_heading(level=token.level).clear())
        self.render_inner(token)

    def render_paragraph(self, token):
        # for all span tokens
        self.paras.append(self.docx.add_paragraph().clear())
        self.render_inner(token)
