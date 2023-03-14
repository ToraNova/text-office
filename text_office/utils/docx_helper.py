'''
Copyright (C) 2023 ToraNova

    This program is free software: you can redistribute it and/or modify
    it under the terms of the GNU Affero General Public License as published
    by the Free Software Foundation, either version 3 of the License, or
    (at your option) any later version.

    This program is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU Affero General Public License for more details.

    You should have received a copy of the GNU Affero General Public License
    along with this program.  If not, see <https://www.gnu.org/licenses/>.
'''

import re
import lxml
from .errx_helper import (
        ensure_valid_attr,
        ensure_valid_value,
        ensure_and_set,
        ensure_all_attr,
        )

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxcompose.composer import Composer
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, CT_P
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.numbering import CT_Num, CT_Numbering, CT_NumPr
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.base import XmlMappedEnumMember, EnumValue
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import (
        Length, Pt, Inches, Cm, Mm, Twips,
        RGBColor,
        )
from docx.styles.styles import Styles
from docx.oxml.xmlchemy import serialize_for_reading
from .parsers import (
        parse_sizespec,
        parse_color,
        parse_sec_orientation,
        parse_table_align,
        parse_para_align,
        parse_bool,
        )

def insert_section(doc, stype):
    vmap = {
        'oddpage': WD_SECTION.ODD_PAGE,
        'evenpage': WD_SECTION.EVEN_PAGE,
        'newpage': WD_SECTION.NEW_PAGE,
        'newcol': WD_SECTION.NEW_COLUMN,
        'cont': WD_SECTION.CONTINUOUS,
        None: WD_SECTION.NEW_PAGE,
    }
    ensure_valid_value('secbr', vmap, stype)
    return doc.add_section(vmap[stype])

def merge_table_cells(table, **kwargs):
    _optmap = {
            'from_row': (int, 'digits'),
            'from_col': (int, 'digits'),
            'to_row': (int, 'digits'),
            'to_col': (int, 'digits'),
            }
    ensure_all_attr(_optmap, kwargs)
    fr = kwargs.get('from_row')
    fc = kwargs.get('from_col')
    tr = kwargs.get('to_row')
    tc = kwargs.get('to_col')

    if tr < fr:
        fr, tr = tr, fr

    if tc < fc:
        fc, tc = tc, fc

    merge_tar = table.cell(fr, fc)
    merge_tar.merge(table.cell(tr, tc))
    return table

def format_section(section, **kwargs):
    _optmap = {
            'left_margin':  ([Length, float], kwargs, section, None, parse_sizespec),
            'right_margin':  ([Length, float], kwargs, section, None, parse_sizespec),
            'top_margin':  ([Length, float], kwargs, section, None, parse_sizespec),
            'bottom_margin':  ([Length, float], kwargs, section, None, parse_sizespec),
            'page_width':  ([Length, float], kwargs, section, None, parse_sizespec),
            'page_height':  ([Length, float], kwargs, section, None, parse_sizespec),
            'orientation': (EnumValue, kwargs, section, None, parse_sec_orientation),
            'header_distance': ([Length, float], kwargs, section, None, parse_sizespec),
            'footer_distance': ([Length, float], kwargs, section, None, parse_sizespec),
            'header_linked': (bool, kwargs, section.header, 'is_linked_to_previous', parse_bool),
            'footer_linked': (bool, kwargs, section.footer, 'is_linked_to_previous', parse_bool)
            }
    ensure_valid_attr(_optmap.keys(), kwargs.keys())

    for k, v in _optmap.items():
        if v is None:
            # special cases
            continue
        ensure_and_set(k, *v)

    # force orientation change by swapping page_width and page_height
    page_orient = parse_sec_orientation(kwargs.get('orientation'))
    if isinstance(page_orient, EnumValue):
        # force orientation change
        pw = section.page_width
        section.page_width = section.page_height
        section.page_height = pw

    return section

def format_run(run, **kwargs):
    _optmap = {
            'bold':         (bool, kwargs, run.font, None, parse_bool),
            'italic':       (bool, kwargs, run.font, None, parse_bool),
            'underline':    (bool, kwargs, run.font, None, parse_bool),
            'strike':       (bool, kwargs, run.font, None, parse_bool),
            'name':         (str, kwargs, run.font),
            'style':        (str, kwargs, run),
            'size':         ([Length, float], kwargs, run.font, None, parse_sizespec),
            'color':        (RGBColor, kwargs, run.font.color, 'rgb', parse_color),
            }
    ensure_valid_attr(_optmap.keys(), kwargs.keys())

    for k, v in _optmap.items():
        if v is None:
            # special cases
            continue
        ensure_and_set(k, *v)
    return run

def format_paragraph(para, **kwargs):
    _optmap = {
            'style': (str, kwargs, para.style),
            'align': (EnumValue, kwargs, para.paragraph_format, 'alignment', parse_para_align),
            'spacing': ([Length, float], kwargs, para.paragraph_format, 'line_spacing', parse_sizespec),
            'before': ([Length, float], kwargs, para.paragraph_format, 'space_before', parse_sizespec),
            'after': ([Length, float], kwargs, para.paragraph_format, 'space_after', parse_sizespec),
            'left_indent': ([Length, float], kwargs, para.paragraph_format, None, parse_sizespec),
            'right_indent': ([Length, float], kwargs, para.paragraph_format, None, parse_sizespec),
            'first_line_indent': ([Length, float], kwargs, para.paragraph_format, None, parse_sizespec),
            'tabstops': None,
            }
    ensure_valid_attr(_optmap.keys(), kwargs.keys())

    for k, v in _optmap.items():
        if v is None:
            # special cases
            continue
        ensure_and_set(k, *v)

    tabstops = kwargs.get('tabstops')
    if isinstance(tabstops, list):
        for ats in tabstops:
            para.paragraph_format.tab_stops.add_tab_stop(*ats)

    return para

def format_figure(figobj, **kwargs):
    _optmap = ('width', 'height', 'border_width', 'border_color')
    ensure_valid_attr(_optmap, kwargs.keys())
    set_figure_dims(figobj, kwargs.get('width'), kwargs.get('height'))

    bw = kwargs.get('border_width')
    bc = kwargs.get('border_color')
    if bw is not None:
        if bc is not None:
            set_figure_border(figobj, bw, bc)
        else:
            set_figure_border(figobj, bw)

def format_table(table, **kwargs):
    _optmap = {
            'style': (str, kwargs, table),
            'align': (EnumValue, kwargs, table, 'alignment', parse_table_align),
            'autofit': (bool, kwargs, table, None, parse_bool),
            'column_widths': None,
            'left_indent': None,
            }
    ensure_valid_attr(_optmap.keys(), kwargs.keys())

    for k, v in _optmap.items():
        if v is None:
            # special cases
            continue
        ensure_and_set(k, *v)

    colwidths = kwargs.get('column_widths')
    if isinstance(colwidths, str):
        cwidths=[]
        for cr in colwidths.split(','):
            cwidths.append(parse_sizespec(cr.strip()))

        table.autofit = False
        for r in table.rows:
            for idx, cell in enumerate(r.cells):
                if idx < len(cwidths):
                    cell.width = cwidths[idx]

    left_indent = parse_sizespec(kwargs.get('left_indent'))
    if left_indent is not None:
        indent_table(table, left_indent)

    return table

def insert_hrule(para, linestyle='single'):
    p = para._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
            )
    bottom = OxmlElement('w:bottom')
    vmap = {
            'dashsmall': 'dashSmallGap',
            'single': 'single',
            None: 'single',
            }
    ensure_valid_value('hr', vmap, linestyle)
    bottom.set(qn('w:val'), vmap[linestyle])
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    return pBdr

def format_tabstops(tabstop, cs):
    if tabstop is None:
        return None

    lmar = Inches(cs.left_margin.inches)
    cmar = Inches(cs.page_width.inches/2 - cs.left_margin.inches)
    rmar = Inches(cs.page_width.inches - (cs.left_margin.inches + cs.right_margin.inches))
    vmap = {
        'left': (lmar, WD_TAB_ALIGNMENT.LEFT),
        'center': (cmar, WD_TAB_ALIGNMENT.CENTER),
        'right': (rmar, WD_TAB_ALIGNMENT.RIGHT),
    }
    out = []
    for t in tabstop.split(','):
        ensure_valid_value('tabstops', vmap, t)
        out.append(vmap[t])
    return out

def insert_hyperlink(para, txt, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = para.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = txt
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = para.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    #r.font.color.rgb = RGBColor(0x06, 0x45, 0xad)
    r.font.underline = True
    return hyperlink

def assign_numbering(doc, para, anid, start=1, tierlvl=0):
    '''
    how to find abstract_numId (anid), open the .docx and find nav to word/document.xml
    search the relevant keyword in the list that we want to restart (i.e. ListBullet)
    open word/styles.xml and search for the style (ListBullet), look for numId value
    open word/numbering.xml and search the numId="x" value, get abstractNumId value
    '''
    ctn = doc.part.numbering_part.numbering_definitions._numbering
    nxtnid = ctn._next_numId
    num = CT_Num.new(nxtnid, anid)
    num.add_lvlOverride(ilvl=tierlvl).add_startOverride(start)
    ctn._insert_num(num)
    para._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = nxtnid
    return nxtnid

def set_paranumpr(para, numid=3, ilvl=0):
    npr = para._p.get_or_add_pPr().get_or_add_numPr()
    npr.get_or_add_numId().val = numid
    npr.get_or_add_ilvl().val = ilvl

def make_caption(run, include_heading=0, caption_type='Figure'):
    '''
    make the run a reference caption. set include_heading=1 to do smth like
    Figure x-y, where the x is the heading numbering to follow, and y the n-th figure under heading x
    '''
    r = run._r

    _strseq = f' SEQ {caption_type} \\* ARABIC '
    if isinstance(include_heading, int) and include_heading > 0:
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = f' STYLEREF {include_heading} \s '
        r.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)

        nbh = OxmlElement('w:noBreakHyphen')
        r.append(nbh)

        _strseq = f' SEQ Figure \\* ARABIC \s {include_heading}'

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = _strseq
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

def delete_paragraph(para):
    p = para._element
    p.getparent().remove(p)
    p._p = p._element = None

def format_cell(cell, **kwargs):
    if 'color' in kwargs:
        color = parse_color(kwargs.pop('color'))
        sh_elem = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(sh_elem)

    if 'align' in kwargs:
        align = kwargs.pop('align')
        for para in cell.paragraphs:
            ikwargs = {'align': align}
            format_paragraph(para, **ikwargs)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    set_border(tcBorders, **kwargs)

    return cell

def set_figure_dims(figobj, width, height):
    if width is None:
        width = figobj.width
    else:
        width = parse_sizespec(width)

    if height is None:
        # maintain aspect ratio here
        sf = float(width) / float(figobj.width)
        figobj.height = round(figobj.height * sf)
    else:
        # set height
        figobj.height = parse_sizespec(height)

    figobj.width = width

def set_figure_border(figobj, border_width='1.5pt', border_color='000000'):

    strw = str(parse_sizespec(border_width))
    colw = str(parse_color(border_color))
    aln = OxmlElement('a:ln')
    aln.set('w', strw)
    afl = OxmlElement('a:solidFill')
    acl = OxmlElement('a:srgbClr')
    acl.set('val', colw)
    afl.append(acl)
    aln.append(afl)

    inline = figobj._inline
    spPr = inline.graphic.graphicData.pic.spPr
    spPr.append(aln)


    extEff = OxmlElement('wp:effectExtent')
    extEff.set('l', strw)
    extEff.set('t', strw)
    extEff.set('r', strw)
    extEff.set('b', strw)

    ext = inline.extent
    ext.addnext(extEff)


def concat_docx(files):
    if isinstance(files, list) and len(files) > 0:
        main_c = Composer(Document(files[0]))

        for f in files[1:]:
            main_c.append(Document(f))

        return main_c.doc
    raise TypeError('input to concat_docx should be a list')

def set_updatefields(docx, val="true"):
    # https://github.com/elapouya/python-docx-template/issues/151#issuecomment-442722594
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    # add child to doc.settings element
    element_updatefields = lxml.etree.SubElement(
        docx.settings.element, f"{namespace}updateFields"
    )
    element_updatefields.set(f"{namespace}val", val)

def insert_TOC(run, prompt_update):
    _add_lox(run, f'TOC \\o "1-3" \\h \\z \\u', prompt_update)

def insert_LOT(run, prompt_update):
    _add_lox(run, f'TOC \\h \\z \\c "Table"', prompt_update)

def insert_LOF(run, prompt_update=True):
    _add_lox(run, f'TOC \\h \\z \\c "Figure"', prompt_update)

def _add_lox(run, ist, prompt_update):
    # https://github.com/xiaominzhaoparadigm/python-docx-add-list-of-tables-figures/blob/master/LOT.py
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    # uncomment the following line to prompt user to update the field
    if prompt_update:
        fldChar.set(qn('w:dirty'), 'true')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ist
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar4)

def insert_pagenum(section, run, **kwargs):
    _optmap = ('show_total', 'start')
    ensure_valid_attr(_optmap, kwargs.keys())

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    if parse_bool(kwargs.get('show_total')):
        # allow 'page x of n' displays
        t2 = OxmlElement('w:t')
        t2.set(qn('xml:space'), 'preserve')
        t2.text = ' of '
        run._r.append(t2)

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "NUMPAGES"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    pstart = kwargs.get('start')
    if isinstance(pstart, str) and pstart.isnumeric():
        pgNumType =  OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), pstart)

        section._sectPr.append(pgNumType)


def format_table_border(table, **kwargs):
    tblPr = table._tblPr

    # check for tag existence, if none found, then create one
    tblBorders = tblPr.first_child_found_in('w:tblBorders')
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    set_border(tblBorders, **kwargs)
    return table


def indent_table(table, indent):
    # noinspection PyProtectedMember
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent.twips))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)


def set_border(tbem, **kwargs):
    _edgemap = ('left', 'right', 'top', 'bottom', 'insideH', 'insideV')
    _typemap = ('width', 'line', 'color', 'space', 'shadow')
    _defaults = {
            'width': '12',
            'color': '000000',
            'line': 'single',
            'space': '0',
            'shadow': 'false',
            }

    _optmap = []
    for e in _edgemap:
        for t in _typemap:
            _optmap.append('%s_%s' % (e, t))

    ensure_valid_attr(_optmap, kwargs.keys())
    submap = {
        'width': 'sz',
        'line': 'val',
    }

    # search for edge options, if any edge options exists, all types should be configured
    for e in _edgemap:
        edge_undefined = True
        for opt in kwargs.keys():
            if opt.startswith(e):
                # one of the types on this edge is defined
                edge_undefined = False
                break

        if edge_undefined:
            # skip the rest
            continue

        for t in _typemap:
            _key = '%s_%s' % (e, t)

            if _key in kwargs:
                edat = kwargs.pop(_key)
            else:
                edat = _defaults[t]

            if t == 'width':
                if edat.endswith('pt'):
                    edat = edat[:-2]

            if t == 'color' and edat != 'auto':
                edat = str(parse_color(edat))
                if not edat.startswith('#'):
                    edat = '#%s' % edat

            if t in submap:
                t = submap[t]

            # check tag existence, if does not exist, create one
            tag = 'w:%s' % e
            element = tbem.find(qn(tag))

            if element is None:
                # element does not exist, create one
                element = OxmlElement(tag)
                tbem.append(element)

            element.set(qn('w:%s' % t), edat)

    return tbem

def left_indent_from_level(obj, lvl, indent_p_lvl=Inches(0.4)):
    if isinstance(obj, Paragraph):
        obj.paragraph_format.left_indent = indent_p_lvl * lvl
    elif isinstance(obj, Table):
        indent_table(obj, Twips(indent_p_lvl.twips * lvl))
