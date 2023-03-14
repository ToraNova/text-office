'''
Copyright (C) 2023 ToraNova

    This program is free software: you can redistribute it and/or modify
    it under the terms of the GNU Affero General Public License as published
    by the Free Software Foundation, either version 3 of the License, or
    (at your option) any later version.

    This program is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU Affero General Public License for more details.

    You should have received a copy of the GNU Affero General Public License
    along with this program.  If not, see <https://www.gnu.org/licenses/>.
'''

# renders a docx from markdown
# using mistletoe for markdown parsing and python-docx for docx rendering

import os
import re
import webcolors

from itertools import chain
from mistletoe import block_token, span_token
from mistletoe.base_renderer import BaseRenderer

from docx import Document
from .. import utils
from ..utils.docx_helper import (
        format_paragraph, format_run, format_table, format_tabstops,
        format_figure, format_cell, format_section, insert_pagenum, format_table_border,
        insert_hyperlink, assign_numbering, insert_section, insert_hrule, merge_table_cells,
        make_caption, delete_paragraph, insert_LOF, insert_LOT, insert_TOC, left_indent_from_level,
        )

from .format_tag import (
        HorizontalRuleTag, LineBreakTag, PageBreakTag, SectionBreakTag, SectionControlTag,
        BoldTag, ItalicTag, UnderlineTag, StrongTag, EmphasisTag, StrikethroughTag,
        FontTag, ImageTag, CellTag, InsertTabTag, InsertPageNumTag, MergeTag,
        TOCTag, LOTTag, LOFTag,
        )

from .format_tag import (
        CommentBlockTag, AlignBlockTag, TableBlockTag, ParagraphBlockTag,
        FooterBlockTag, HeaderBlockTag, BorderBlockTag,
        )

class Renderer(BaseRenderer):
    def __init__(self, tag='notag', docx_template=None, rel_root=None, docx_opts=None, *extras):
        '''
        renders a .docx using docx_template as the template file
        rel_root determine the path to include relative path resources (e.g., images)
        '''
        self._suppress_ptag_stack = [False]
        super().__init__(*chain(([
            HorizontalRuleTag, LineBreakTag, PageBreakTag,
            SectionBreakTag, InsertTabTag, SectionControlTag,
            BoldTag, ItalicTag, UnderlineTag, StrongTag, EmphasisTag, StrikethroughTag,
            FontTag, ImageTag, CellTag, InsertPageNumTag, MergeTag,
            CommentBlockTag, AlignBlockTag, TableBlockTag, ParagraphBlockTag,
            FooterBlockTag, HeaderBlockTag, BorderBlockTag,
            TOCTag, LOTTag, LOFTag,
        ]), extras))
        self.rel_root = os.getcwd() if rel_root is None else rel_root
        self.docx_template = docx_template
        self.docx_opts = {}
        self.run_tag = tag
        if isinstance(docx_opts, dict):
            self.docx_opts = docx_opts

    def render_inner(self, token):
        # inner rendering not supported on DocxRenderer
        # raise NotImplementedError('unsupported render_inner function called')
        for c in token.children:
            self.render(c)

    def render_list(self, token):
        # https://github.com/miyuchina/mistletoe/blob/94022647cd9d80e242db5c93a6567e3155b468bc/mistletoe/block_token.py#L447
        ordered = token.start is not None
        level = self.list_level
        self.list_level += 1

        #if 'List Number' in self.docx.styles and 'List Bullet' in self.docx.styles:
        # old numbering system

        # TODO: fix this abstract_numId map with a more robust implementation
        anid_map = {
            0: 7, # numId 5, absnumId 7
            1: 3, # numId 6, absnumId 3
            2: 2, # numId 7, absnumId 2
        }

        if ordered:
            stpl = 'List Number '
        else:
            stpl = 'List Bullet '

        if level > 0:
            stpl += str(level+1)
        style = stpl.strip()

        for idx, c in enumerate(token.children):
            number = token.start + idx if ordered else None
            # if mistletoe changed the way leader/prepend is computed, this must changed as well

            tos = len(self.paras)
            for ic in c.children:
                ic.docx_style = style
                self.render(ic)
            #self.populate_and_format_paras(c, style=style)

            added = len(self.paras) - tos
            #if added > 0 and ordered:
            if added > 0 and ordered and idx == 0:
                assign_numbering(self.docx, self.paras[tos], anid_map[level], start=number)

        #elif 'List Paragraph' in self.docx.styles:

        #    nid_map = {
        #        True: 3, # ordered use numId = 3
        #        False: 5 # non-ordered use numId = 5, check the .docx document.xml
        #    }

        #    for idx, c in enumerate(token.children):
        #        number = token.start + idx if ordered else 1

        #        tos = len(self.paras)
        #        for ic in c.children:
        #            ic.docx_style = 'List Paragraph'
        #            self.render(ic)
        #        added = len(self.paras) - tos
        #        set_paranumpr(self.paras[tos], nid_map[ordered], level)
        #else:
        #    raise Exception('no valid numbering styles found')

        self.list_level -= 1

    def render_list_item(self, token):
        # handled by list
        #self.render_inner(token)
        raise NotImplementedError("list_item is handled by list rendering")

    def render_link(self, token):
        # TODO: allow alt-txt in links, see mistletoe docs on parser issues (token.title is empty!)
        if len(token.title) < 1:
            ut = token.target
        else:
            ut = token.title
        insert_hyperlink(self.paras[-1], ut, token.target)

    def render_auto_link(self, token):
        insert_hyperlink(self.paras[-1], token.target, token.target)

    def render_image(self, token):
        if token.src.startswith('/'):
            img_path = token.src
        else:
            img_path = os.path.join(self.rel_root, token.src)

        if not os.path.isfile(img_path):
            raise FileNotFoundError(f'image file \'{img_path}\' not found')

        self.runs.append(self.paras[-1].add_run())
        self.inline_shapes.append(self.runs[-1].add_picture(img_path))

        # skip caption if no caption string
        if len(token.title.strip()) < 1:
            return

        self.populate_caption('Figure', token.title)

    def render_block_code(self, token):
        _dxopt = utils.parse_bool(self.docx_opts.get('error_on_blockcode', False))
        if _dxopt:
            raise NotImplementedError('block_codes not supported, use <font name="Lucida Sans Typewriter"></font> instead')
        else:
            # just render as-is on ms-docx
            self.render_paragraph(token)

    def render_document(self, token):
        # create document from template
        if self.docx_template is None:
            self.docx = Document()
        else:
            self.docx = Document(self.docx_template)
            delete_paragraph(self.docx.paragraphs[-1])

        # paragraph and run stack
        self.list_level = 0 # default root level
        self.heading_level = None # no heading level
        self.runs = []
        self.paras = self.docx.paragraphs.copy()
        self.tables = self.docx.tables.copy()
        self.sections = []
        for s in self.docx.sections:
            self.sections.append(s)
        self.cells = []
        self.inline_shapes = []
        self.render_inner(token)

        # save document
        #self.docx.save(self.out_path)
        return self.docx

    def populate_and_format_runs(self, token, **kwargs):
        # format_tag allows nested formatting <b><i>test</i></b>, italic and bold
        tos = len(self.runs) # top of stack
        self.render_inner(token)
        added = len(self.runs) - tos # new top of stack
        # apply format to all added elements during inner render
        for i in range(added):
            format_run(self.runs[-(i+1)], **kwargs)

    def populate_and_format_paras(self, token, **kwargs):
        tos = len(self.paras)
        self.render_inner(token)
        added = len(self.paras) - tos
        # apply format to all added elements during inner render
        for i in range(added):
            tar = self.paras[-(i+1)]
            format_paragraph(tar, **kwargs)

    def populate_and_format_tables(self, token, **kwargs):
        tos = len(self.tables) # top of stack
        self.render_inner(token)
        added = len(self.tables) - tos # new top of stack
        # add style to all added table within this block
        for i in range(added):
            format_table(self.tables[-(i+1)], **kwargs)

    def render_font_tag(self, token):
        self.populate_and_format_runs(token, **token.format)

    def render_bold_tag(self, token):
        self.populate_and_format_runs(token, bold=True)

    def render_strong(self, token):
        self.populate_and_format_runs(token, style='Strong', bold=True)

    def render_strong_tag(self, token):
        self.populate_and_format_runs(token, style='Strong')

    def render_italic_tag(self, token):
        self.populate_and_format_runs(token, italic=True)

    def render_emphasis(self, token):
        self.populate_and_format_runs(token, style='Emphasis', italic=True)

    def render_emphasis_tag(self, token):
        self.populate_and_format_runs(token, style='Emphasis')

    def render_underline_tag(self, token):
        self.populate_and_format_runs(token, underline=True)

    def render_strikethrough_tag(self, token):
        self.populate_and_format_runs(token, strike=True)

    def render_strikethrough(self, token):
        self.populate_and_format_runs(token, strike=True)

    def render_heading(self, token):
        # assume that heading has no additional child
        self.heading_level = token.level
        tar = self.docx.add_heading(level=token.level).clear()
        self.paras.append(tar)
        self.render_inner(token)
        self.auto_left_indent(tar, token.level-1)

    def render_paragraph(self, token):
        # for all span tokens

        render_ptr = self.docx
        if hasattr(token, 'render_to') and isinstance(token.render_to, str):
            if token.render_to == 'header':
                render_ptr = self.sections[-1].header

            elif token.render_to == 'footer':
                render_ptr = self.sections[-1].footer
        else:
            token.render_to = 'body'

        if hasattr(token, 'docx_style') and isinstance(token.docx_style, str):
            par = render_ptr.add_paragraph(style=token.docx_style)
        else:
            par = render_ptr.add_paragraph()
        par.clear()

        self.paras.append(par)
        self.render_inner(token)
        self.auto_left_indent(par)

    def auto_left_indent(self, obj, lvl = None):
        _dxopt = utils.parse_docx_sizespec(self.docx_opts.get('auto_left_indent', None))
        if not _dxopt:
            return
        if not isinstance(self.heading_level, int):
            # no headings yet, left most
            lvl = 0
        elif not isinstance(lvl, int):
            # there are headings, but lvl is not an int (default)
            lvl = self.heading_level
        left_indent_from_level(obj, lvl, _dxopt)

    def render_horizontal_rule_tag(self, token):
        insert_hrule(self.paras[-1], token.format_value)

    def render_thematic_break(self, token):
        insert_hrule(self.paras[-1])

    def render_page_break_tag(self, token):
        self.docx.add_page_break()

    def render_toc_tag(self, token):
        self.runs.append(self.paras[-1].add_run())
        _dxopt = utils.parse_bool(self.docx_opts.get('prompt_updatefield', True))
        insert_TOC(self.runs[-1], _dxopt)

    def render_lot_tag(self, token):
        self.runs.append(self.paras[-1].add_run())
        _dxopt = utils.parse_bool(self.docx_opts.get('prompt_updatefield', True))
        insert_LOT(self.runs[-1], _dxopt)

    def render_lof_tag(self, token):
        self.runs.append(self.paras[-1].add_run())
        _dxopt = utils.parse_bool(self.docx_opts.get('prompt_updatefield', True))
        insert_LOF(self.runs[-1], _dxopt)

    def render_raw_text(self, token):
        # add run to last added paragraph
        self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_text(token.content)

    def render_line_break(self, token):
        self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_break()

    def render_line_break_tag(self, token):
        self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_break()

    def render_insert_tab_tag(self, token):
        # insert a tab character
        if len(self.runs) < 1:
            self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_tab()

    def render_insert_page_num_tag(self, token):
        self.runs.append(self.paras[-1].add_run())
        insert_pagenum(self.sections[-1], self.runs[-1], **token.format)

    def render_section_break_tag(self, token):
        # insert new section
        new_section = insert_section(self.docx, token.format_value)
        new_section.header.is_linked_to_previous = False
        new_section.footer.is_linked_to_previous = False
        self.sections.append(new_section)

    def render_section_control_tag(self, token):
        # make adjustments to current section
        format_section(self.sections[-1], **token.format)

    def render_comment_block_tag(self, token):
        # do nothing
        pass

    def render_table(self, token):
        if len(token.children) < 1:
            return
        ncol = len(token.children[0].children)
        tbltar = self.docx.add_table(0, ncol)
        self.tables.append(tbltar)

        if hasattr(token, 'header') and token.header is not None:
            token.header.is_header = True
            self.render(token.header)

        for row in token.children:
            row.is_header = False
            self.render(row)

        self.auto_left_indent(tbltar)


    def render_table_row(self, token):
        row = self.tables[-1].add_row()
        for cidx, col in enumerate(token.children):
            self.cells.append(row.cells[cidx])
            self.render(col)

    def render_table_cell(self, token):
        self.paras.append(self.cells[-1].paragraphs[0])  # add default para on cell
        self.render_inner(token)

    def render_table_block_tag(self, token):
        # additional valid opts
        if isinstance(token.format.get('caption'), str):
            kwargs = {
                'align': token.format.pop('caption_align', None)
            }
            self.populate_caption('Table', token.format.get('caption'), **kwargs)
            token.format.pop('caption')
        self.populate_and_format_tables(token, **token.format)

    def render_cell_tag(self, token):
        if len(self.cells) < 1:
            return

        self.render_inner(token)
        format_cell(self.cells[-1], **token.format)

    def render_merge_tag(self, token):
        if len(self.tables) < 1:
            return

        merge_table_cells(self.tables[-1], **token.format)

    def render_border_block_tag(self, token):
        tos = len(self.tables)
        self.render_inner(token)
        added = len(self.tables) - tos
        for i in range(added):
            target = self.tables[-(i+1)]
            format_table_border(target, **token.format)

    def render_align_block_tag(self, token):
        self.populate_and_format_paras(token, align=token.format_value)

    def populate_caption(self, caption_type, caption_string, **kwargs):
        tcpar = self.docx.add_paragraph(style='Caption').clear()
        self.paras.append(tcpar)
        # self.auto_left_indent(tcpar)  # uncomment this to auto-indent captions
        format_paragraph(tcpar, **kwargs)
        _dxopt = int(self.docx_opts.get('caption_prefix_heading', 0))
        make_caption(tcpar.add_run(f'{caption_type} '), _dxopt, caption_type)

        tcrun = tcpar.add_run(f': {caption_string}')

    def render_image_tag(self, token):
        tos = len(self.inline_shapes)
        self.render_inner(token)
        added = len(self.inline_shapes) - tos
        for i in range(added):
            target = self.inline_shapes[-(i+1)]
            format_figure(target, **token.format)

    def render_paragraph_block_tag(self, token):
        self.populate_and_format_paras(token, **token.format)

    def render_header_block_tag(self, token):
        _tsalign = format_tabstops(token.format.get('tabstops'), self.sections[-1])
        utils.set_attr_recursively(token, block_token.Paragraph, 'render_to', 'header')
        self.populate_and_format_paras(token, tabstops=_tsalign)

    def render_footer_block_tag(self, token):
        _tsalign = format_tabstops(token.format.get('tabstops'), self.sections[-1])
        utils.set_attr_recursively(token, block_token.Paragraph, 'render_to', 'footer')
        self.populate_and_format_paras(token, tabstops=_tsalign)
